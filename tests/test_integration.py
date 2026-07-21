"""集成测试"""
import os
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

import pytest


class TestProfileManager:
    def setup_method(self):
        from config import PROFILE_PATH
        if PROFILE_PATH.exists():
            PROFILE_PATH.unlink()

    def test_create_and_load(self):
        from tools.profile_manager import create_empty_profile, save_profile, load_profile
        p = create_empty_profile()
        assert "meta" in p
        assert len(p["fields"]) == 18
        assert p["fields"]["name"]["status"] == "missing"

    def test_check_completeness(self):
        from tools.profile_manager import create_empty_profile, check_completeness
        p = create_empty_profile()
        assert len(check_completeness(p)) == 9

    def test_merge_extracted_fields(self):
        from tools.profile_manager import merge_extracted_fields
        data = {
            "name": "张三", "career_objective": ["Python开发"],
            "expected_salary": "20K", "location": "深圳",
            "phone": "13800138000", "email": "test@test.com",
            "education": [{"school": "XX大学", "major": "CS", "degree": "本科", "graduation_year": "2020"}],
            "skills": [{"name": "Python", "level": "精通"}],
            "projects": [{"name": "项目A", "role": "开发", "start_date": "2021", "end_date": "2022", "description": "test"}],
            "personal_strengths": "优点", "self_evaluation": "自我评价",
        }
        profile = merge_extracted_fields(data, ["test.pdf"])
        assert profile["fields"]["name"]["value"] == "张三"
        assert profile["fields"]["email"]["value"] == "test@test.com"


class TestResumeParser:
    def test_parse_txt(self):
        from tools.resume_parser import parse_file
        with tempfile.NamedTemporaryFile(suffix=".txt", mode="w", encoding="utf-8", delete=False) as f:
            f.write("姓名：张三\n技能：Python, Java")
            path = f.name
        text = parse_file(path)
        assert "张三" in text
        assert "Python" in text
        os.unlink(path)

    def test_parse_docx(self):
        from tools.resume_parser import parse_file
        from docx import Document
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_paragraph("姓名：李四")
            doc.add_paragraph("技能：Django, React")
            doc.save(f.name)
            path = f.name
        text = parse_file(path)
        assert "李四" in text
        assert "Django" in text
        os.unlink(path)

    def test_unsupported_format(self):
        from tools.resume_parser import parse_file
        with pytest.raises(ValueError, match="Unsupported"):
            parse_file("test.xyz")


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
