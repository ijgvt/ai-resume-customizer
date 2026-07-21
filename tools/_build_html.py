#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
简历优化工具：读取简历文件 + 生成 HTML 简历

用法:
  python generate_resume.py read <file_path>          读取 PDF/Word/TXT 简历，输出纯文本 JSON
  python generate_resume.py build <json_file> <output> 读取结构化 JSON，生成 HTML 简历文件
"""

import json
import os
import sys
import base64
import html as html_mod


def configure_output_encoding():
    for stream in (sys.stdout, sys.stderr):
        reconfigure = getattr(stream, "reconfigure", None)
        if callable(reconfigure):
            try:
                reconfigure(encoding="utf-8", errors="replace")
            except Exception:
                pass


configure_output_encoding()


# ---------------------------------------------------------------------------
# 简历读取
# ---------------------------------------------------------------------------

def read_pdf(file_path: str) -> str:
    try:
        import fitz
        doc = fitz.open(file_path)
        parts = []
        for page in doc:
            parts.append(page.get_text())
        doc.close()
        return "\n".join(parts).strip()
    except ImportError:
        return ""
    except Exception:
        return ""


def read_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)
        return "\n".join(paragraphs).strip()
    except ImportError:
        return ""
    except Exception:
        return ""


def read_txt(file_path: str) -> str:
    for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            with open(file_path, "r", encoding=enc) as f:
                return f.read().strip()
        except (UnicodeDecodeError, UnicodeError):
            continue
    return ""


def read_resume(file_path: str) -> dict:
    if not os.path.isfile(file_path):
        return {"error": f"文件不存在: {file_path}"}
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        text = read_pdf(file_path)
    elif ext == ".docx":
        text = read_docx(file_path)
    elif ext in (".txt", ".text", ".md"):
        text = read_txt(file_path)
    else:
        return {"error": f"不支持的文件格式: {ext}，支持 PDF / Word(.docx) / TXT"}
    if not text:
        return {"error": "文件内容为空或无法解析，请尝试复制简历文本直接粘贴"}
    return {
        "status": "success",
        "file_path": file_path,
        "file_type": ext,
        "text": text,
        "char_count": len(text),
    }


# ---------------------------------------------------------------------------
# HTML 简历生成 — 参考产品经理简历样式
# ---------------------------------------------------------------------------

HTML_TEMPLATE = r"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{name} - 简历</title>
<style>
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{
  font-family: "PingFang SC", "Microsoft YaHei", "Source Han Sans SC", sans-serif;
  color: #333;
  line-height: 1.65;
  background: #f0f0f0;
  font-size: 13px;
}}
.resume {{
  max-width: 210mm;
  margin: 20px auto;
  padding: 18mm 20mm 14mm;
  background: white;
  box-shadow: 0 1px 6px rgba(0,0,0,0.08);
}}
/* ---- Header ---- */
.header {{
  display: flex;
  justify-content: space-between;
  align-items: flex-start;
  padding-bottom: 8px;
  border-bottom: 1.5px solid #2c3e50;
  margin-bottom: 12px;
}}
.header-left h1 {{
  font-size: 22px;
  font-weight: 700;
  color: #1a1a1a;
  letter-spacing: 1px;
  margin-bottom: 4px;
}}
.header-contact {{
  font-size: 12px;
  color: #555;
  line-height: 1.7;
}}
.header-contact .label {{
  color: #888;
}}
.header-photo {{
  width: 108px;
  height: 144px;
  object-fit: cover;
  border-radius: 4px;
  border: 1px solid #ddd;
  flex-shrink: 0;
}}
/* ---- Section ---- */
.section {{
  margin-bottom: 10px;
}}
.section-title {{
  font-size: 14px;
  font-weight: 700;
  color: #1a1a1a;
  border-bottom: 1px solid #ddd;
  padding-bottom: 3px;
  margin-bottom: 6px;
}}
/* ---- Education ---- */
.edu-row {{
  display: flex;
  justify-content: space-between;
  align-items: baseline;
  margin-bottom: 1px;
  font-size: 13px;
}}
.edu-left {{ display: flex; align-items: baseline; gap: 6px; }}
.edu-school {{ font-weight: 700; }}
.edu-tag {{
  display: inline-block;
  background: #e8f4fd;
  color: #2980b9;
  font-size: 10.5px;
  padding: 0 5px;
  border-radius: 2px;
  font-weight: 600;
}}
.edu-detail {{ color: #555; }}
.edu-date {{ color: #888; font-size: 12px; white-space: nowrap; }}
.edu-courses {{
  font-size: 12px;
  color: #666;
  margin-top: 1px;
  line-height: 1.6;
}}
/* ---- Experience ---- */
.exp-header {{
  display: flex;
  justify-content: space-between;
  align-items: baseline;
  margin-bottom: 4px;
}}
.exp-left {{ display: flex; align-items: baseline; gap: 8px; }}
.exp-company {{ font-weight: 700; font-size: 13px; }}
.exp-role {{ color: #555; font-size: 13px; }}
.exp-date {{ color: #888; font-size: 12px; white-space: nowrap; }}
.bullet-list {{
  margin: 0;
  padding-left: 0;
  list-style: none;
}}
.bullet-list li {{
  position: relative;
  padding-left: 12px;
  margin-bottom: 2px;
  font-size: 12.5px;
  line-height: 1.65;
  color: #444;
}}
.bullet-list li::before {{
  content: "\00b7";
  position: absolute;
  left: 1px;
  color: #888;
  font-weight: bold;
}}
.bullet-label {{
  font-weight: 600;
  color: #222;
}}
/* ---- Skills ---- */
.skill-item {{
  margin-bottom: 2px;
  font-size: 12.5px;
  line-height: 1.65;
  padding-left: 12px;
  position: relative;
  color: #444;
}}
.skill-item::before {{
  content: "\00b7";
  position: absolute;
  left: 1px;
  color: #888;
  font-weight: bold;
}}
.skill-label {{
  font-weight: 600;
  color: #222;
}}
/* ---- Self Evaluation ---- */
.self-evaluation {{
  font-size: 12.5px;
  line-height: 1.65;
  color: #444;
}}
/* ---- Print ---- */
@page {{
  size: A4;
  margin: 0;
}}
@media print {{
  body {{ background: white; }}
  .resume {{
    margin: 0; padding: 18mm 20mm 14mm; box-shadow: none; max-width: none;
    transform: none !important;
    width: auto !important;
    transform-origin: unset !important;
  }}
  .section {{ page-break-inside: avoid; }}
}}
</style>
</head>
<body>
<div class="resume" id="resume">
  {header_html}
  {education_html}
  {experience_html}
  {skills_html}
  {self_evaluation_html}
</div>
<script>
(function() {{
  var el = document.getElementById("resume");
  if (!el) return;
  var pageH = 297 - 18 - 14;
  if (el.scrollHeight > pageH * 3.78) {{
    el.style.transformOrigin = "top left";
    var scale = (pageH * 3.78) / el.scrollHeight;
    el.style.transform = "scale(" + scale.toFixed(4) + ")";
    el.style.width = (100 / scale).toFixed(2) + "%";
  }}
}})();
</script>
</body>
</html>"""


def _esc(text: str) -> str:
    return html_mod.escape(str(text)) if text else ""


def _encode_photo_base64(photo_path: str, max_height: int = 288) -> str:
    """读取照片文件，调整高度后返回 base64 data URI。"""
    if not photo_path or not os.path.isfile(photo_path):
        return ""
    ext = os.path.splitext(photo_path)[1].lower()
    mime_map = {".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".png": "image/png"}
    mime = mime_map.get(ext, "")
    if not mime:
        return ""
    try:
        with open(photo_path, "rb") as f:
            raw = f.read()
        try:
            from PIL import Image
            import io
            img = Image.open(io.BytesIO(raw))
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            w, h = img.size
            if h > max_height:
                ratio = max_height / h
                img = img.resize((int(w * ratio), max_height), Image.LANCZOS)
            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=92)
            b64 = base64.b64encode(buf.getvalue()).decode("ascii")
            mime = "image/jpeg"
        except ImportError:
            b64 = base64.b64encode(raw).decode("ascii")
        return f"data:{mime};base64,{b64}"
    except Exception:
        return ""


def _build_header_html(data: dict) -> str:
    name = _esc(data.get("name", ""))
    phone = _esc(data.get("phone", ""))
    email = _esc(data.get("email", ""))
    objective = _esc(data.get("objective", ""))
    photo_path = data.get("photo", "")
    contact_lines = []
    if phone:
        contact_lines.append(f'<span class="label">电话：</span>{phone}')
    if email:
        contact_lines.append(f'<span class="label">邮箱：</span>{email}')
    if objective:
        contact_lines.append(f'<span class="label">求职意向：</span>{objective}')
    contact_html = "<br>".join(contact_lines)
    photo_b64 = _encode_photo_base64(photo_path) if photo_path else ""
    photo_tag = f'\n  <img class="header-photo" src="{photo_b64}" alt="证件照">' if photo_b64 else ""
    return (
        '<div class="header">\n'
        f'  <div class="header-left"><h1>{name}</h1>\n'
        f'  <div class="header-contact">{contact_html}</div></div>\n'
        f'  {photo_tag}\n'
        "</div>"
    )


def _build_education_html(edu: dict) -> str:
    if not edu:
        return ""
    school = _esc(edu.get("school", ""))
    tag = _esc(edu.get("school_tag", ""))
    major = _esc(edu.get("major", ""))
    degree = _esc(edu.get("degree", ""))
    period = _esc(edu.get("period", ""))
    gpa = edu.get("gpa", "")
    courses = edu.get("courses", [])
    tag_html = f'<span class="edu-tag">{tag}</span>' if tag else ""
    detail_parts = [p for p in [major, degree] if p]
    detail = " | ".join(detail_parts)
    gpa_suffix = f" | GPA: {_esc(gpa)}" if gpa else ""
    courses_html = ""
    if courses:
        courses_text = "、".join(_esc(c) for c in courses)
        courses_html = f'<div class="edu-courses">● 主修课程：{courses_text}</div>'
    return (
        '<div class="section">\n'
        '  <div class="section-title">教育背景</div>\n'
        f'  <div class="edu-row">\n'
        f'    <div class="edu-left"><span class="edu-school">{school}</span>{tag_html}<span class="edu-detail">&nbsp;|&nbsp;{detail}{gpa_suffix}</span></div>\n'
        f'    <span class="edu-date">{period}</span>\n'
        f"  </div>\n"
        f"  {courses_html}\n"
        "</div>"
    )


def _build_experience_html(experience: list) -> str:
    if not experience:
        return ""
    sections = {}
    for item in experience:
        sec = item.get("section", "经历")
        sections.setdefault(sec, []).append(item)
    html_parts = []
    for sec_name, items in sections.items():
        html_parts.append(f'<div class="section">\n  <div class="section-title">{_esc(sec_name)}</div>')
        for item in items:
            company = _esc(item.get("company", ""))
            role = _esc(item.get("role", "") or item.get("title", ""))
            period = _esc(item.get("period", ""))
            bullets = item.get("bullets", [])
            # 兼容旧格式: situation/task/action/result
            if not bullets:
                for key, label in [("situation", "场景背景"), ("task", "任务职责"), ("action", "行动方法"), ("result", "项目成果")]:
                    val = item.get(key, "")
                    if val:
                        bullets.append({"label": label, "text": val})
            bullet_html = ""
            for b in bullets:
                bl = _esc(b.get("label", ""))
                bt = _esc(b.get("text", ""))
                if bl and bt:
                    bullet_html += f'<li><span class="bullet-label">{bl}：</span>{bt}</li>\n'
                elif bt:
                    bullet_html += f'<li>{bt}</li>\n'
            html_parts.append(
                f'  <div class="exp-header">\n'
                f'    <div class="exp-left"><span class="exp-company">{company}</span><span class="exp-role">{role}</span></div>\n'
                f'    <span class="exp-date">{period}</span>\n'
                f"  </div>\n"
                f'  <ul class="bullet-list">\n{bullet_html}  </ul>'
            )
        html_parts.append("</div>")
    return "\n".join(html_parts)


def _build_skills_html(skills: list) -> str:
    if not skills:
        return ""
    html_parts = ['<div class="section">\n  <div class="section-title">岗位技能</div>']
    for s in skills:
        label = _esc(s.get("label", ""))
        text = _esc(s.get("text", ""))
        if label and text:
            html_parts.append(f'  <div class="skill-item"><span class="skill-label">{label}：</span>{text}</div>')
        elif text:
            html_parts.append(f'  <div class="skill-item">{text}</div>')
    html_parts.append("</div>")
    return "\n".join(html_parts)


def _build_self_evaluation_html(text: str) -> str:
    if not text:
        return ""
    return (
        '<div class="section">\n'
        '  <div class="section-title">自我评价</div>\n'
        f'  <div class="self-evaluation">{_esc(text)}</div>\n'
        '</div>'
    )


def generate_html(resume_data: dict) -> str:
    header_html = _build_header_html(resume_data)
    education_html = _build_education_html(resume_data.get("education", {}))
    experience_html = _build_experience_html(resume_data.get("experience", []))
    skills_html = _build_skills_html(resume_data.get("skills", []))
    self_evaluation_html = _build_self_evaluation_html(resume_data.get("self_evaluation", ""))
    return HTML_TEMPLATE.format(
        name=_esc(resume_data.get("name", "简历")),
        header_html=header_html,
        education_html=education_html,
        experience_html=experience_html,
        skills_html=skills_html,
        self_evaluation_html=self_evaluation_html,
    )


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> int:
    if len(sys.argv) < 2:
        print(json.dumps({"error": "用法: python generate_resume.py read|build ..."}, ensure_ascii=False))
        return 1
    command = sys.argv[1]
    if command == "read":
        if len(sys.argv) < 3:
            print(json.dumps({"error": "请提供简历文件路径"}, ensure_ascii=False))
            return 1
        result = read_resume(sys.argv[2])
        print(json.dumps(result, ensure_ascii=False, indent=2))
        return 0 if result.get("status") == "success" else 1
    elif command == "build":
        if len(sys.argv) < 4:
            print(json.dumps({"error": "用法: python generate_resume.py build <json_file> <output_path>"}, ensure_ascii=False))
            return 1
        json_path = sys.argv[2]
        output_path = sys.argv[3]
        if not os.path.isfile(json_path):
            print(json.dumps({"error": f"JSON 文件不存在: {json_path}"}, ensure_ascii=False))
            return 1
        with open(json_path, "r", encoding="utf-8") as f:
            resume_data = json.load(f)
        html_content = generate_html(resume_data)
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html_content)
        print(json.dumps({
            "status": "success",
            "output_path": os.path.abspath(output_path),
            "file_size": len(html_content),
        }, ensure_ascii=False, indent=2))
        return 0
    else:
        print(json.dumps({"error": f"未知命令: {command}，支持 read 或 build"}, ensure_ascii=False))
        return 1


if __name__ == "__main__":
    sys.exit(main())
